"""Export 12 scénarios shorts Coco — formats variés, ancrés dans du réel."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "SCENARIOS_v2.docx"

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

t = doc.add_heading("Scénarios shorts — Coco the Axolotl", level=0)
for run in t.runs:
    run.font.color.rgb = RGBColor(0xC2, 0x3B, 0x7E)

intro = doc.add_paragraph()
intro.add_run(
    "12 scénarios répartis en 4 formats. Tous ancrés dans du réel (biologie, "
    "confession auteur, aesthetic, réactions enfants) — pas de hooks « j'ai testé X "
    "pendant 6 semaines ». Format 9:16, 20-45s. Difficulté de production indiquée."
).italic = True
doc.add_paragraph()

doc.add_heading("Formats utilisés", level=1)
doc.add_paragraph("• Format A — Curiosité axolotl (5 scénarios) : un fait biologique réel mène au livre", style="List Bullet")
doc.add_paragraph("• Format B — Auteur / coulisses (3 scénarios) : confession, time-lapse, origin story", style="List Bullet")
doc.add_paragraph("• Format C — Page reveal esthétique (2 scénarios) : ASMR, une seule ligne du livre, ambiance", style="List Bullet")
doc.add_paragraph("• Format D — Réactions enfants brutes (2 scénarios) : caméra sur l'enfant, zéro voiceover marketing", style="List Bullet")
doc.add_paragraph()


# =============== SCENARIOS ===============
scenarios = [
    # ---------- FORMAT A — Curiosité axolotl ----------
    {
        "section": "A — Curiosité axolotl (faits biologiques réels → livre)",
        "items": [
            {
                "n": "A1", "tag": "neoteny / never grow up",
                "title": "Axolotls never grow up. So I wrote about one.",
                "diff": "Faible — texte plein écran + 1 page livre",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Texte plein écran fond noir, typo manuscrite : « Axolotls never grow up. »"),
                    ("Hook 2", "3-6s", "(silence)", "Macro shot d'un axolotl rose réel (stock / Fal.ai), gros plan branchies. Overlay : « It's called neoteny. »"),
                    ("Setup", "6-13s", "They stay babies their whole lives. Pink. Smiling. Forever soft.", "B-roll axolotl en aquarium, mouvement lent"),
                    ("Reveal", "13-22s", "I made a children's book about one who can't sleep. Because the part of us that stays a kid — never sleeps either.", "Livre Coco Can't Sleep tenu, ouvert sur l'illustration de Coco les yeux ouverts"),
                    ("CTA", "22-28s", "Coco the Axolotl. On Amazon.", "Couverture livre + lien"),
                ],
            },
            {
                "n": "A2", "tag": "regeneration / love",
                "title": "Axolotls regenerate their own heart.",
                "diff": "Faible — texte + page livre",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Texte fond noir : « Axolotls regenerate. »"),
                    ("Hook 2", "3-6s", "(silence)", "Macro axolotl. Overlay : « Even their heart. »"),
                    ("Setup", "6-13s", "Cut their leg, it grows back. Damage their heart, it heals. They're the only animal we know who does that.", "Stock axolotl"),
                    ("Reveal", "13-22s", "Maybe that's why I wrote a book about an axolotl mom who loves harder than seems possible. Hearts that can rebuild can love bigger.", "Livre I Love You More, page maman/Coco"),
                    ("CTA", "22-28s", "I Love You More. Coco the Axolotl.", "CTA"),
                ],
            },
            {
                "n": "A3", "tag": "permanent smile",
                "title": "Why axolotls always look like they're smiling.",
                "diff": "Faible — fact + brand teaser",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Macro axolotl qui « sourit ». Overlay : « Axolotls smile. »"),
                    ("Hook 2", "3-6s", "(silence)", "Même plan. Overlay : « They literally can't stop. »"),
                    ("Setup", "6-13s", "Their face muscles can't form any other expression. The smile is permanent. Anatomy.", "Plans rapprochés visage axolotl"),
                    ("Reveal", "13-22s", "I built a book brand around that. A character who can only be soft. For kids who need a face that won't ever frown back at them.", "Livre Coco couverture + dos + 3 livres alignés"),
                    ("CTA", "22-28s", "Coco the Axolotl. Three books on Amazon.", "CTA"),
                ],
            },
            {
                "n": "A4", "tag": "transparent eggs",
                "title": "You can watch axolotl babies hatch through the shell.",
                "diff": "Moyenne — stock vidéo œuf transparent",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Macro œuf translucide axolotl, embryon visible. Overlay : « You can SEE them hatch. »"),
                    ("Hook 2", "3-6s", "(silence)", "Embryon qui bouge dans l'œuf. Overlay : « Through the shell. »"),
                    ("Setup", "6-13s", "Their eggs are translucent. You watch the baby form, day by day, like an aquarium of becoming.", "Stock time-lapse œuf axolotl"),
                    ("Reveal", "13-22s", "That's the book I wrote. Coco finds eggs in the meadow and asks: whose is this? The reader watches the answer hatch.", "Livre Whose Egg Is This?, page mystère + page reveal"),
                    ("CTA", "22-28s", "Whose Egg Is This. On Amazon.", "CTA"),
                ],
            },
            {
                "n": "A5", "tag": "extinction",
                "title": "There are 1000 wild axolotls left. So I'm making them famous.",
                "diff": "Faible — texte + livre",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Texte : « Less than 1000 left. »"),
                    ("Hook 2", "3-6s", "(silence)", "Texte : « In the wild. »"),
                    ("Setup", "6-15s", "Axolotls live in one lake outside Mexico City. Pollution, invasive fish, urbanization. They're vanishing.", "Stock lac Xochimilco / axolotl rare"),
                    ("Reveal", "15-25s", "I can't save them. But I can put one in millions of kids' hands. So they grow up knowing this animal exists. Coco the Axolotl.", "Trois livres alignés"),
                    ("CTA", "25-32s", "Three books. On Amazon. A few seconds of fame for a vanishing species.", "CTA"),
                ],
            },
        ],
    },
    # ---------- FORMAT B — Auteur / coulisses ----------
    {
        "section": "B — Auteur / coulisses (confession, time-lapse, origin)",
        "items": [
            {
                "n": "B1", "tag": "time-lapse drawing",
                "title": "I drew Coco's face 47 times before I got it right.",
                "diff": "Moyenne — il faut filmer un time-lapse de dessin",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Time-lapse main qui dessine, 1ère version maladroite. Overlay : « Try 1. »"),
                    ("Hook 2", "3-6s", "(silence)", "Versions 2-10 défilent rapidement. Overlay : « Try 47 was the one. »"),
                    ("Setup", "6-15s", "I needed his eyes to feel safe to a 4-year-old at 9pm. Half-closed. Soft. Not sleepy enough to seem dead. Not awake enough to be wired.", "Time-lapse zoom sur les yeux"),
                    ("Reveal", "15-25s", "Forty-seven tries. This is what landed.", "Couverture finale Coco Can't Sleep, gros plan visage"),
                    ("CTA", "25-32s", "Coco the Axolotl. Worth the 47 tries. Amazon link in bio.", "CTA"),
                ],
            },
            {
                "n": "B2", "tag": "the line that took 6 months",
                "title": "The line in my kids book that took 6 months to write.",
                "diff": "Faible — texte + voix",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Carnet avec ratures, beaucoup de versions barrées. Overlay : « 6 months on one line. »"),
                    ("Hook 2", "3-6s", "(silence)", "Stylo qui pose enfin une phrase. Overlay : « It had to be exact. »"),
                    ("Setup", "6-15s", "The line my own mom said when I was scared at night. I wanted to give it to other kids. But every version sounded fake.", "Pages de carnet manuscrit"),
                    ("Reveal", "15-28s", "The line I landed on : « The dolphin will come. He always does. » That's it. Simple. True. Repeats like a heartbeat.", "Page du livre où la phrase apparaît, voix qui lit la phrase"),
                    ("CTA", "28-35s", "Coco Can't Sleep Tonight. Amazon.", "CTA"),
                ],
            },
            {
                "n": "B3", "tag": "origin story",
                "title": "I wrote a bedtime book for the anxious kid I was.",
                "diff": "Faible — face cam + livre",
                "rows": [
                    ("Hook 1", "0-3s", "(silence)", "Face cam auteur, regard direct, lumière douce. Overlay : « I was a kid who didn't sleep. »"),
                    ("Hook 2", "3-6s", "(silence)", "Même plan. Overlay : « So I wrote this for myself. »"),
                    ("Setup", "6-18s", "Anxious. Catastrophizing. My mom would sit on the bed and tell me a dolphin was coming to take my worries to the sea. Made up. Worked.", "Plans rapides : photo enfance, lit, fenêtre ouverte"),
                    ("Reveal", "18-30s", "Years later I wrote it down so other anxious kids could borrow my mom's voice. That's the book.", "Livre Coco Can't Sleep tenu, ouvert"),
                    ("CTA", "30-38s", "Coco the Axolotl. For the ones who lie awake.", "CTA"),
                ],
            },
        ],
    },
    # ---------- FORMAT C — Page reveal esthétique ----------
    {
        "section": "C — Page reveal esthétique (ASMR, une ligne, ambiance)",
        "items": [
            {
                "n": "C1", "tag": "ASMR page turn",
                "title": "A page from Coco Can't Sleep Tonight.",
                "diff": "Faible — filmer une page tournée + voix douce",
                "rows": [
                    ("0-2s", "0-2s", "(silence, son ambient)", "Plan macro main qui ouvre le livre, son du papier, ambient drone très bas"),
                    ("2-15s", "2-15s", "« Coco watched the moon. The moon watched him back. They didn't say anything. They didn't need to. »", "Plan fixe sur la page, voix off masculine grave très lente, son ambient eau/vent"),
                    ("15-22s", "15-22s", "(silence)", "La main tourne la page lentement. Couverture du livre apparaît en surimpression coin"),
                    ("CTA", "22-28s", "(voix douce) Coco the Axolotl.", "Logo + lien discret bas d'écran"),
                ],
            },
            {
                "n": "C2", "tag": "ASMR I Love You More",
                "title": "A line my mom wishes she'd had.",
                "diff": "Faible — page filmée + voix",
                "rows": [
                    ("0-3s", "0-3s", "(silence)", "Texte sur écran noir : « A line my mom wishes she'd had. »"),
                    ("3-18s", "3-18s", "« I love you more than the polar bear loves her cub through the longest winter. »", "Plan macro page polar bear du livre, voix féminine lente, son ambient feu de cheminée"),
                    ("18-25s", "18-25s", "(silence)", "Page tournée, illustration éléphant apparaît"),
                    ("CTA", "25-30s", "(voix) I Love You More. Coco the Axolotl.", "CTA discret"),
                ],
            },
        ],
    },
    # ---------- FORMAT D — Réactions enfants brutes ----------
    {
        "section": "D — Réactions enfants brutes (caméra sur l'enfant, zéro voix marketing)",
        "items": [
            {
                "n": "D1", "tag": "live reaction",
                "title": "She didn't expect the dolphin.",
                "diff": "Moyenne — il faut filmer un vrai enfant",
                "rows": [
                    ("0-3s", "0-3s", "(silence)", "Caméra sur visage enfant, livre invisible, parent qui lit hors champ. Overlay : « Watch her face. »"),
                    ("3-15s", "3-15s", "(audio diégétique : voix parent qui lit) « Coco looked out his window. And then... »", "Caméra reste sur l'enfant. Yeux qui s'agrandissent"),
                    ("15-22s", "15-22s", "(audio enfant : « A DOLPHIN ?? »)", "L'enfant se retourne vers le livre, bouche ouverte"),
                    ("22-30s", "22-30s", "(silence, juste sourire enfant)", "Cut sur livre dans ses mains, elle tourne les pages elle-même. Overlay : « Coco Can't Sleep Tonight »"),
                    ("CTA", "30-35s", "(silence)", "Logo Coco + Amazon link"),
                ],
            },
            {
                "n": "D2", "tag": "again",
                "title": "Request count : 47.",
                "diff": "Moyenne — montage 47 voix d'enfants disant 'again' (ou 1 enfant 47 prises)",
                "rows": [
                    ("0-3s", "0-3s", "(audio compilé : « again », « again », « encore », « otra vez »)", "Cuts ultra rapides : différents enfants qui réclament le livre, mains qui tendent l'objet vers parent"),
                    ("3-10s", "3-10s", "(continue audio « again »)", "Compilation voix se densifie, overlay compteur qui monte : « 1 », « 12 », « 47 »"),
                    ("10-18s", "10-18s", "(silence)", "Plan large : enfant assis dans son lit, livre ouvert sur ses genoux"),
                    ("18-25s", "18-25s", "(voix off une seule fois) « Coco the Axolotl. Most-requested book in our house. »", "Couverture livre"),
                    ("CTA", "25-32s", "(silence)", "CTA Amazon + lien"),
                ],
            },
        ],
    },
]


def add_scenario(item):
    h = doc.add_heading(f"{item['n']} — {item['tag']}", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x80, 0x2D, 0x60)

    p = doc.add_paragraph()
    r = p.add_run("Title : ")
    r.bold = True
    p.add_run(item["title"])

    p = doc.add_paragraph()
    r = p.add_run("Difficulté : ")
    r.bold = True
    p.add_run(item["diff"]).italic = True

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 4"
    hdr = table.rows[0].cells
    hdr[0].text = "Beat"
    hdr[1].text = "Time"
    hdr[2].text = "Voice / audio"
    hdr[3].text = "Visual"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for beat, t_, voice, visual in item["rows"]:
        row = table.add_row().cells
        row[0].text = beat
        row[1].text = t_
        row[2].text = voice
        row[3].text = visual
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()


for block in scenarios:
    h = doc.add_heading(block["section"], level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xC2, 0x3B, 0x7E)
    for it in block["items"]:
        add_scenario(it)
    doc.add_paragraph()


doc.add_heading("Reco — par où commencer", level=1)
doc.add_paragraph(
    "Tester d'abord les scénarios à difficulté FAIBLE (texte + page livre) pour mesurer "
    "ce qui résonne sans investir en tournage :"
)
doc.add_paragraph("• A1 — Axolotls never grow up (curiosité brute, hook universel)", style="List Bullet")
doc.add_paragraph("• B2 — La phrase qui m'a pris 6 mois (auteur authentique)", style="List Bullet")
doc.add_paragraph("• C1 — ASMR page reveal (aesthetic, partageable)", style="List Bullet")
doc.add_paragraph()
doc.add_paragraph(
    "Si un format performe, on en produit 5-10 variations sur le même angle. "
    "Si aucun ne performe en 2 semaines, on pivot."
).italic = True


doc.save(OUT)
print(f"OK -> {OUT}")
