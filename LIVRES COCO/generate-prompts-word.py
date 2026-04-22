from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('ChatGPT / DALL-E Prompts — "Mommy Loves You More"', 0)
title.runs[0].font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)

# Important note
doc.add_heading('IMPORTANT : Upload your Coco reference image first!', level=2)
p = doc.add_paragraph(
    'Before using any prompt below, upload your reference image of Coco to ChatGPT and say:'
)
doc.add_paragraph(
    '"Here is my character Coco. Use this exact character design for all illustrations I ask you to '
    'create. Keep the style consistent throughout."'
).runs[0].italic = True

doc.add_paragraph()

# Style anchor
doc.add_heading('STYLE ANCHOR (à inclure dans TOUS les prompts)', level=2)
style_text = (
    "STYLE: Cute children's book illustration, soft watercolor style, pastel colors, rounded shapes, "
    "warm and friendly atmosphere. NO text in the illustration.\n\n"
    "Coco is a small pink axolotl walking upright on two legs, with a wide happy smile, big round dark "
    "eyes, smooth soft pink body, small round pink ears/gills on each side of the head, cartoon-style "
    "hands with EXACTLY 4 fingers (like Mickey Mouse — not 5), wearing cute pink sneakers.\n\n"
    "Mama is a taller pink axolotl, same species as Coco, with a warm and loving expression. "
    "She walks upright. Her hands also have exactly 4 fingers. "
    "She wears a soft, flowing dress and cute ballet flats (NOT sneakers, never sneakers).\n\n"
    "Square format, 1:1 ratio."
)
p = doc.add_paragraph(style_text)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)
p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# All prompts
prompts = [
    {
        "title": "COVER",
        "text": (
            "Children's book cover illustration. Nighttime magical scene under a sky full of golden stars.\n\n"
            "Mama axolotl (taller, pink, warm smile) stands on the ground and holds little Coco (small pink axolotl, pink sneakers) "
            "in a warm, protective hug. Coco is nestled against Mama's chest, completely safe and held. "
            "Coco's face shows pure happiness and security. Mama's arms wrap all the way around Coco.\n\n"
            "Both characters are standing/grounded. NO floating. The hug feels like the safest place in the world.\n\n"
            "Around them: a warm golden glow. Stars and sparkles everywhere. The night sky is deep blue and violet.\n\n"
            "Leave space at the TOP for the title text and BOTTOM for the author name.\n\n"
            "Square format 1:1. Watercolor children's book style, soft pastel colors, magical and full of love."
        )
    },
    {
        "title": "p.1 — Bedtime / Mama tucks Coco in",
        "text": (
            "Children's book illustration. Cozy bedroom scene at night.\n\n"
            "Mama axolotl (tall, pink, loving smile) gently tucks little Coco (small, pink, "
            "sleepy happy face) into a warm cozy bed. The blanket is pulled up just right — not too tight, "
            "not too loose.\n\n"
            "Soft moonlight through the window. A small nightlight glows. Stuffed animal on the pillow next "
            "to Coco. Everything feels safe and warm.\n\n"
            "Square format 1:1. Soft blues and pinks, watercolor children's book style. Maximum coziness."
        )
    },
    {
        "title": "p.2 — First \"I love you\" exchange",
        "text": (
            "Children's book illustration. Intimate nighttime bedroom scene.\n\n"
            "Close-up: Little Coco (small pink axolotl) lies in bed looking up at Mama with "
            "big round loving eyes, full of warmth. Mama axolotl leans in close, face near Coco's, whispering "
            "with the most tender smile.\n\n"
            "Soft moonlight. The moment feels very close, very quiet, very safe.\n\n"
            "Square format 1:1. Soft blues and warm pinks, watercolor style. Intimate and emotional."
        )
    },
    {
        "title": "p.3 — Coco wonders",
        "text": (
            "Children's book illustration. Cozy bedroom at night.\n\n"
            "Coco (small pink axolotl) sits up in bed with one finger on chin, thinking very "
            "hard. Big curious eyes. A tiny smile on the face.\n\n"
            "Around Coco's head: small playful thought bubbles floating — question marks, stars, tiny hearts. "
            "The expression is adorable and funny.\n\n"
            "Square format 1:1. Soft pastel blues and pinks, watercolor style. Playful and curious."
        )
    },
    {
        "title": "p.4 — Morning hug (Coco hugs Mama)",
        "text": (
            "Children's book illustration. Warm sunny kitchen in the morning.\n\n"
            "Coco (small pink axolotl, pink sneakers) hugs Mama axolotl tight from behind, arms "
            "wrapped all the way around her waist, face pressed against her back with a huge smile. Mama "
            "looks over her shoulder laughing.\n\n"
            "Breakfast on the table: pancakes, orange juice, a little bowl. Golden morning light streams "
            "through the window.\n\n"
            "Square format 1:1. Warm yellows and pinks, watercolor style. Joyful and cozy."
        )
    },
    {
        "title": "p.5 — Mama hugs back even tighter",
        "text": (
            "Children's book illustration. Warm sunny kitchen.\n\n"
            "Mama axolotl turns and hugs little Coco back in a HUGE hug — arms completely wrapped around "
            "Coco, squeezing tight. Both are laughing with pure joy. Coco's feet barely touch the floor. "
            "The hug is enormous and warm.\n\n"
            "Square format 1:1. Warm golden tones, watercolor style. Pure joy and laughter."
        )
    },
    {
        "title": "p.6 — Park: Coco stretches arms wide",
        "text": (
            "Children's book illustration. Sunny park on a bright day.\n\n"
            "Coco (small pink axolotl, pink sneakers) stands in the middle of a green park, "
            "arms stretched out as wide as they can possibly go — as wide as the whole world. Big grin, "
            "eyes sparkling with excitement.\n\n"
            "Flowers in the background, blue sky, a few white clouds. Everything is bright and joyful.\n\n"
            "Square format 1:1. Bright greens and yellows, watercolor style. Energetic and joyful."
        )
    },
    {
        "title": "p.7 — Mama stretches even wider",
        "text": (
            "Children's book illustration. Sunny park.\n\n"
            "Mama axolotl stretches her arms out even wider than Coco — so wide they seem to go beyond the "
            "edges of the page. She's laughing out loud. Little Coco looks up at Mama's arms in total "
            "amazement and starts to giggle.\n\n"
            "The scale is funny: Mama's arms are comically enormous compared to Coco.\n\n"
            "Square format 1:1. Bright park colors, watercolor style. Funny and playful."
        )
    },
    {
        "title": "p.8 — Night: Coco points to the stars",
        "text": (
            "Children's book illustration. Beautiful night outside.\n\n"
            "Coco (small pink axolotl, pink sneakers) and Mama sit together on the grass under a "
            "spectacular starry sky. Coco points up at the stars with one little arm, face full of wonder. "
            "Mama looks up too, smiling.\n\n"
            "The sky is deep blue and violet, filled with hundreds of bright stars and a crescent moon. "
            "A gentle warm light surrounds them.\n\n"
            "Square format 1:1. Deep blues and silvers with warm golden accents, watercolor style. Magical and wonder-filled."
        )
    },
    {
        "title": "p.9 — Mama points past the stars",
        "text": (
            "Children's book illustration. Magical night sky.\n\n"
            "Mama axolotl points dramatically upward — past the stars, past the moon, past everything "
            "visible — into the infinite deep space. Her expression is certain and full of love. Little "
            "Coco follows her gaze, completely in awe, mouth slightly open.\n\n"
            "Above them: stars, moon, galaxies fading into infinite darkness.\n\n"
            "Square format 1:1. Deep blues, silvers and soft purples, watercolor style. Dramatic and awe-inspiring."
        )
    },
    {
        "title": "p.10 — Coco asks the question",
        "text": (
            "Children's book illustration. Quiet nighttime scene outside.\n\n"
            "Close-up: Coco (small pink axolotl) looks up at Mama with the most genuinely "
            "curious and innocent expression. Head slightly tilted. Eyes wide and honest.\n\n"
            "Mama looks down at Coco with infinite warmth. Stars sparkle softly above them. The moment is "
            "quiet and intimate.\n\n"
            "Square format 1:1. Soft blues and warm pinks, watercolor style. Quiet and emotional."
        )
    },
    {
        "title": "p.11 — Mama pulls Coco close",
        "text": (
            "Children's book illustration. Peaceful night outside under the stars.\n\n"
            "Mama axolotl sits down on the grass and gently pulls little Coco into her arms, holding them "
            "close. Both are quiet. Coco's head rests against Mama. Stars above them. A warm golden glow "
            "surrounds them both.\n\n"
            "The scene is still. Tender. The most peaceful image.\n\n"
            "Square format 1:1. Deep blues and warm gold, watercolor style. Tender and emotional."
        )
    },
    {
        "title": "p.12 — Before you were born (most emotional)",
        "text": (
            "Children's book illustration. Dreamlike, magical scene.\n\n"
            "Mama axolotl holds a tiny baby Coco in her arms for the very first time. The baby is impossibly "
            "small, eyes just barely open, curled up against Mama's chest. Mama's face shows pure overwhelming "
            "love and awe.\n\n"
            "Soft golden light radiates around them. Stars drift gently around the scene like snowflakes. "
            "Everything feels magical, warm and sacred.\n\n"
            "This is the most emotional illustration in the book.\n\n"
            "Square format 1:1. Deep warm blues and soft gold, dreamlike watercolor style. Maximum emotional impact."
        )
    },
    {
        "title": "p.13 — Three vignettes: first smile, first word, first hug",
        "text": (
            "Children's book illustration. Dreamlike, magical atmosphere.\n\n"
            "Three small glowing circular vignettes arranged like stars in a dark sky:\n\n"
            "1. Baby Coco smiling for the very first time — tiny, radiant smile\n"
            "2. Baby Coco opening their little mouth saying their first word, eyes wide with surprise\n"
            "3. Baby Coco and Mama in their very first hug — tiny arms wrapped around her\n\n"
            "Each vignette glows with warm golden light. The background is deep night blue with soft stars "
            "connecting them.\n\n"
            "Square format 1:1. Deep blue with warm gold accents, watercolor style. Tender and nostalgic."
        )
    },
    {
        "title": "p.14 — I will always love you more",
        "text": (
            "Children's book illustration. Warm intimate scene.\n\n"
            "Mama axolotl kneels down to Coco's level, holding Coco's small face gently in both hands, "
            "looking directly into Coco's eyes. Mama's expression is calm, certain, full of infinite love "
            "— this is the truest thing she has ever said.\n\n"
            "Coco looks back at Mama, completely still, completely held.\n\n"
            "Warm golden light between them. Soft stars behind.\n\n"
            "Square format 1:1. Warm golds and pinks, watercolor style. Close-up, intimate, the most loving image."
        )
    },
    {
        "title": "p.15 — Coco thinks for a long time",
        "text": (
            "Children's book illustration. Quiet night scene.\n\n"
            "Coco (small pink axolotl) sits alone under the stars, very still, thinking very "
            "deeply. Eyes wide open, a small thoughtful smile beginning to form. Stars twinkle gently above.\n\n"
            "The whole scene feels like a quiet moment before something wonderful happens.\n\n"
            "Square format 1:1. Soft blues and pinks, watercolor style. Quiet, anticipatory, slightly magical."
        )
    },
    {
        "title": "p.16 — The biggest smile",
        "text": (
            "Children's book illustration. Bright joyful scene.\n\n"
            "Close-up of Coco (small pink axolotl) with THE BIGGEST, most radiant, most joyful "
            "smile possible. Eyes sparkling with light. Cheeks round with happiness. One bright star shines "
            "directly above Coco's head.\n\n"
            "This is the happiest Coco has ever been.\n\n"
            "Square format 1:1. Bright warm pinks and golds, watercolor style. Pure, radiant joy."
        )
    },
    {
        "title": "p.17 — Coco wins! (triumphant)",
        "text": (
            "Children's book illustration. Playful joyful scene.\n\n"
            "Coco (small pink axolotl, pink sneakers) points one little finger triumphantly at Mama, "
            "laughing with complete delight. Eyes full of mischief and love. Coco looks like they just won "
            "the best game in the world.\n\n"
            "Mama axolotl looks completely delighted and surprised — she did NOT see this coming. She laughs too.\n\n"
            "Square format 1:1. Warm bright pinks, watercolor style. Funny, triumphant, full of love."
        )
    },
    {
        "title": "p.18 — Okay. You win.",
        "text": (
            "Children's book illustration. Joyful scene.\n\n"
            "Mama axolotl laughs out loud — head back, eyes crinkled, pure joy. Happy tears in her eyes. "
            "She looks completely won over and delighted.\n\n"
            "Little Coco jumps and dances nearby with absolute glee — arms in the air, sneakers off the ground.\n\n"
            "Both are glowing with laughter and happiness.\n\n"
            "Square format 1:1. Warm bright pinks and yellows, watercolor style. Joyful, funny, full of laughter."
        )
    },
    {
        "title": "p.19 — The big hug (finale)",
        "text": (
            "Children's book illustration. FINAL SCENE — the most beautiful image in the book.\n\n"
            "Mama axolotl and Coco are in the biggest, warmest, most perfect hug. Both are laughing and a "
            "little teary with happiness. Stars surround them like a constellation. A warm golden light glows "
            "from between them.\n\n"
            "The night is beautiful. Everything is still. This is the end of the game — and the beginning of forever.\n\n"
            "Square format 1:1. Deep blues and warm gold, watercolor style. The most beautiful, emotional, peaceful image."
        )
    },
    {
        "title": "BACK COVER",
        "text": (
            "Children's book back cover illustration. Peaceful, warm scene.\n\n"
            "Mama axolotl and little Coco (pink sneakers) walk home side by side, hand in hand, "
            "under a sky full of stars. Mama is taller, Coco is small. Both face forward, relaxed and happy.\n\n"
            "A warm golden light glows ahead of them on the path home. Stars above. Flowers along the path.\n\n"
            "Leave space at the BOTTOM for a few lines of text.\n\n"
            "Square format 1:1. Soft blues and warm golds, watercolor style. Peaceful, tender, beautiful."
        )
    },
]

for prompt in prompts:
    doc.add_heading(f'Prompt {prompt["title"]}', level=2)
    p = doc.add_paragraph(prompt["text"])
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    doc.add_paragraph()

output_path = r'C:\Users\33612\Documents\site_COCOtheaxolotl\prototypes\chatgpt-prompts-mommy-loves-you-more-v5.docx'
doc.save(output_path)
print(f"Word document saved: {output_path}")
